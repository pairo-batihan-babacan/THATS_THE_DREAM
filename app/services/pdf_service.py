import io
import os
import sys
import shutil
import subprocess
import tempfile
import zipfile
from pypdf import PdfWriter, PdfReader
import img2pdf
from app.core.config import settings

GS_BINARY = "/usr/bin/gs"


def _output_path(job_id: str, filename: str) -> str:
    os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
    return os.path.join(settings.OUTPUT_DIR, filename)


def compress_pdf(input_path: str, job_id: str, quality: str = "medium") -> str:
    """Compress a PDF using Ghostscript. quality: low | medium | high"""
    output_filename = f"{job_id}_compressed.pdf"
    output_path = _output_path(job_id, output_filename)

    # Map quality level to Ghostscript PDFSETTINGS
    gs_quality = {"low": "/screen", "medium": "/ebook", "high": "/printer"}.get(quality, "/ebook")

    result = subprocess.run(
        [
            GS_BINARY,
            "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.4",
            f"-dPDFSETTINGS={gs_quality}",
            "-dNOPAUSE",
            "-dQUIET",
            "-dBATCH",
            f"-sOutputFile={output_path}",
            input_path,
        ],
        capture_output=True,
        text=True,
    )
    if result.returncode == 0:
        return output_filename

    # Ghostscript failed — fall back to pypdf lossless compression
    gs_error = result.stderr.strip() or result.stdout.strip() or "Unknown Ghostscript error"
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)
        return output_filename
    except Exception as fallback_err:
        raise RuntimeError(
            f"PDF compression failed. "
            f"Ghostscript error: {gs_error}. "
            f"Fallback error: {fallback_err}"
        )


def pdf_to_word(input_path: str, job_id: str) -> str:
    """Convert PDF to DOCX with two-tier approach:
    1. LibreOffice (fast, mature C++ engine, handles most PDFs well)
    2. PyMuPDF text extraction → python-docx (always works, layout-free fallback)
    """
    output_filename = f"{job_id}_converted.docx"
    output_path = _output_path(job_id, output_filename)
    last_error = "Unknown error"

    # Tier 1: LibreOffice PDF→DOCX (same engine used for Word/PPT→PDF, very reliable)
    lo_home = f"/tmp/lo_home_{job_id}"
    os.makedirs(lo_home, exist_ok=True)
    env = os.environ.copy()
    env["HOME"] = lo_home
    try:
        result = subprocess.run(
            [
                "libreoffice", "--headless", "--norestore", "--nofirststartwizard",
                "--convert-to", "docx",
                "--outdir", settings.OUTPUT_DIR,
                input_path,
            ],
            capture_output=True, text=True, timeout=300, env=env,
        )
        # LibreOffice names output after the input filename
        base = os.path.splitext(os.path.basename(input_path))[0]
        lo_output = os.path.join(settings.OUTPUT_DIR, f"{base}.docx")
        if result.returncode == 0 and os.path.exists(lo_output):
            os.rename(lo_output, output_path)
            return output_filename
        last_error = result.stderr.strip() or result.stdout.strip() or "LibreOffice failed"
    except subprocess.TimeoutExpired:
        last_error = "LibreOffice timed out"
    finally:
        shutil.rmtree(lo_home, ignore_errors=True)

    # Tier 2: PyMuPDF text extraction → python-docx (always works, no layout)
    try:
        import fitz
        from docx import Document as DocxDocument
        doc = fitz.open(input_path)
        word = DocxDocument()
        word.add_paragraph(
            "Note: This PDF's layout could not be fully reconstructed. "
            "Text content has been extracted as-is."
        )
        for page in doc:
            text = page.get_text()
            if text.strip():
                word.add_paragraph(text)
            word.add_page_break()
        doc.close()
        word.save(output_path)
        return output_filename
    except Exception as fallback_err:
        raise RuntimeError(
            f"PDF to Word conversion failed: {last_error} "
            f"(text fallback also failed: {fallback_err})"
        )


def pdf_to_images(input_path: str, job_id: str) -> str:
    """Convert PDF pages to JPEG images using PyMuPDF, zip and return.
    PyMuPDF is 5-10x faster than poppler and frees each page from RAM immediately."""
    import fitz  # PyMuPDF

    zip_filename = f"{job_id}_pages.zip"
    zip_path = _output_path(job_id, zip_filename)
    mat = fitz.Matrix(150 / 72, 150 / 72)  # 150 DPI

    try:
        doc = fitz.open(input_path)
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_STORED) as zf:
                for i, page in enumerate(doc):
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img_path = os.path.join(tmp_dir, f"page_{i + 1}.jpg")
                    pix.save(img_path)
                    zf.write(img_path, f"{job_id}_page_{i + 1}.jpg")
                    pix = None  # release pixmap memory immediately
        doc.close()
    except Exception as e:
        raise RuntimeError(
            f"Failed to convert PDF to images. "
            f"Ensure the file is a valid PDF. Error: {e}"
        )

    return zip_filename


def images_to_pdf(input_paths: list[str], job_id: str) -> str:
    output_filename = f"{job_id}_merged.pdf"
    output_path = _output_path(job_id, output_filename)
    try:
        with open(output_path, "wb") as f:
            img2pdf.convert(*input_paths, outputstream=f)
    except Exception as e:
        raise RuntimeError(f"Failed to combine images into PDF: {e}")
    return output_filename


def merge_pdfs(input_paths: list[str], job_id: str) -> str:
    output_filename = f"{job_id}_merged.pdf"
    output_path = _output_path(job_id, output_filename)
    writer = PdfWriter()
    for path in input_paths:
        try:
            reader = PdfReader(path)
            for page in reader.pages:
                writer.add_page(page)
        except Exception as e:
            raise RuntimeError(f"Failed to read PDF '{os.path.basename(path)}': {e}")
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_filename


def split_pdf(input_path: str, job_id: str, pages_spec: str) -> str:
    """
    pages_spec format: '1-3,5,7-9' (1-indexed)
    Returns a zip of the extracted pages as a single PDF.
    """
    try:
        reader = PdfReader(input_path)
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF: {e}")

    page_numbers = _parse_page_spec(pages_spec, len(reader.pages))
    if not page_numbers:
        raise ValueError(
            f"Page spec '{pages_spec}' produced no valid pages. "
            f"The PDF has {len(reader.pages)} page(s)."
        )

    writer = PdfWriter()
    for n in page_numbers:
        writer.add_page(reader.pages[n])

    output_filename = f"{job_id}_split.pdf"
    output_path = _output_path(job_id, output_filename)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_filename


def _parse_page_spec(spec: str, total: int) -> list[int]:
    """Convert '1-3,5,7-9' to zero-indexed page list."""
    pages = []
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-")
            pages.extend(range(int(start) - 1, int(end)))
        else:
            pages.append(int(part) - 1)
    return [p for p in pages if 0 <= p < total]


def strip_pdf_metadata(input_path: str, job_id: str) -> str:
    output_filename = f"{job_id}_clean.pdf"
    output_path = _output_path(job_id, output_filename)
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        # Clear Info dictionary
        writer.add_metadata({})
        # Remove XMP metadata stream from the document catalog
        writer._root_object.pop("/Metadata", None)
        with open(output_path, "wb") as f:
            writer.write(f)
    except Exception as e:
        raise RuntimeError(f"Failed to strip PDF metadata: {e}")
    return output_filename


def inspect_pdf_metadata(content: bytes) -> dict:
    """Return a dict of Info-dict and XMP presence from PDF bytes."""
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF: {e}")

    result = {}
    if reader.metadata:
        for key, value in reader.metadata.items():
            clean_key = key.lstrip("/")
            result[clean_key] = str(value)

    # Check for XMP stream
    try:
        xmp = reader.xmp_metadata
        if xmp:
            result["XMP"] = "Present (XML metadata stream)"
    except Exception:
        pass

    return result


def _libreoffice_to_pdf(input_path: str, job_id: str) -> str:
    """Convert any LibreOffice-compatible document to PDF (shared helper)."""
    output_filename = f"{job_id}_converted.pdf"
    output_dir = settings.OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)

    # Each job gets its own LibreOffice HOME to prevent concurrent profile lock conflicts
    lo_home = f"/tmp/lo_home_{job_id}"
    os.makedirs(lo_home, exist_ok=True)

    env = os.environ.copy()
    env["HOME"] = lo_home

    try:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--norestore",
                "--nofirststartwizard",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                input_path,
            ],
            capture_output=True,
            text=True,
            env=env,
            timeout=120,
        )
    finally:
        shutil.rmtree(lo_home, ignore_errors=True)

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice failed to convert '{os.path.basename(input_path)}' to PDF. "
            f"Error: {result.stderr.strip() or result.stdout.strip() or 'Unknown error'}"
        )

    # LibreOffice names the output after the input filename (without extension)
    base = os.path.splitext(os.path.basename(input_path))[0]
    libreoffice_output = os.path.join(output_dir, f"{base}.pdf")
    final_path = os.path.join(output_dir, output_filename)

    if not os.path.exists(libreoffice_output):
        raise RuntimeError(
            f"LibreOffice reported success but output file was not found "
            f"(expected: {libreoffice_output}). stdout: {result.stdout.strip()}"
        )

    os.rename(libreoffice_output, final_path)
    return output_filename


def word_to_pdf(input_path: str, job_id: str) -> str:
    return _libreoffice_to_pdf(input_path, job_id)


def ppt_to_pdf(input_path: str, job_id: str) -> str:
    """Convert PPT/PPTX presentation to PDF using LibreOffice."""
    return _libreoffice_to_pdf(input_path, job_id)


def excel_to_pdf(input_path: str, job_id: str) -> str:
    """Convert XLS/XLSX spreadsheet to PDF using LibreOffice."""
    return _libreoffice_to_pdf(input_path, job_id)


def unlock_pdf(input_path: str, job_id: str, password: str = "") -> str:
    """Remove password protection from a PDF file."""
    output_filename = f"{job_id}_unlocked.pdf"
    output_path = _output_path(job_id, output_filename)
    try:
        reader = PdfReader(input_path)
        if reader.is_encrypted:
            if not reader.decrypt(password):
                raise RuntimeError("Incorrect password — could not unlock this PDF")
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to unlock PDF: {e}")
    return output_filename


def rotate_pdf(input_path: str, job_id: str, rotations: dict) -> str:
    """Apply per-page rotation deltas. rotations: {"0": 90, "3": 180} (0-indexed, degrees to add)."""
    import fitz

    doc = fitz.open(input_path)
    for page_idx_str, delta in rotations.items():
        idx = int(page_idx_str)
        if 0 <= idx < len(doc):
            page = doc[idx]
            page.set_rotation((page.rotation + int(delta)) % 360)

    output_filename = f"{job_id}_rotated.pdf"
    output_path = _output_path(job_id, output_filename)
    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    return output_filename


def analyze_pdf_orientation(content: bytes) -> list[dict]:
    """Read page dimensions and current rotation from PDF bytes to detect landscape pages."""
    import fitz

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF: {e}")

    pages = []
    for i, page in enumerate(doc):
        r = page.rect
        w, h = round(r.width, 1), round(r.height, 1)
        pages.append({
            "index": i,
            "page_number": i + 1,
            "width": w,
            "height": h,
            "rotation": page.rotation,
            "is_landscape": w > h,
        })
    doc.close()
    return pages


def protect_pdf(input_path: str, job_id: str, password: str, owner_password: str = "") -> str:
    """Encrypt a PDF with AES-256. password is required to open the file."""
    if not password or not password.strip():
        raise ValueError("A password is required to protect a PDF")

    output_filename = f"{job_id}_protected.pdf"
    output_path = _output_path(job_id, output_filename)

    try:
        reader = PdfReader(input_path)
        if reader.is_encrypted:
            raise RuntimeError(
                "This PDF is already password-protected. "
                "Use Unlock PDF to remove the existing protection first, "
                "then re-protect with a new password."
            )
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        # AES-256 — strongest PDF encryption standard
        # user_password: must be entered to open the file
        # owner_password: controls permissions (print/copy/edit); defaults to user_password
        writer.encrypt(
            user_password=password,
            owner_password=owner_password if owner_password else password,
            algorithm="AES-256",
        )
        with open(output_path, "wb") as f:
            writer.write(f)
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to protect PDF: {e}")

    return output_filename


def flatten_pdf(input_path: str, job_id: str) -> str:
    """Flatten form fields and annotations into static page content.

    Uses PyMuPDF show_pdf_page to render each source page (including all widget
    appearances and annotations) as static vector content into a new PDF — no
    AcroForm, no interactivity, text remains searchable.
    Fallback: Ghostscript pdfwrite re-render.
    """
    import fitz

    output_filename = f"{job_id}_flattened.pdf"
    output_path = _output_path(job_id, output_filename)

    try:
        src = fitz.open(input_path)
        out = fitz.open()
        for i, page in enumerate(src):
            # Creates a new page and renders the source page (with all widget appearances
            # baked in as static vector content) — no interactive elements survive
            new_page = out.new_page(width=page.rect.width, height=page.rect.height)
            new_page.show_pdf_page(new_page.rect, src, i)
        src.close()
        out.save(output_path, garbage=4, deflate=True)
        out.close()
        return output_filename
    except Exception as fitz_err:
        # Fallback: Ghostscript pdfwrite re-render
        gs_result = subprocess.run(
            [
                GS_BINARY,
                "-sDEVICE=pdfwrite",
                "-dCompatibilityLevel=1.4",
                "-dNOPAUSE", "-dQUIET", "-dBATCH",
                f"-sOutputFile={output_path}",
                input_path,
            ],
            capture_output=True,
            text=True,
        )
        if gs_result.returncode == 0 and os.path.exists(output_path):
            return output_filename
        raise RuntimeError(
            f"PDF flatten failed. "
            f"PyMuPDF: {fitz_err}. "
            f"Ghostscript: {gs_result.stderr.strip() or 'failed'}"
        )


def pdf_to_excel(input_path: str, job_id: str) -> str:
    """Extract text content from a PDF and export it as an Excel spreadsheet."""
    import pandas as pd  # pandas is already in requirements

    output_filename = f"{job_id}_data.xlsx"
    output_path = _output_path(job_id, output_filename)
    try:
        reader = PdfReader(input_path)
        rows = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            for line in text.splitlines():
                line = line.strip()
                if line:
                    rows.append({"Page": page_num, "Content": line})
        if not rows:
            rows = [{"Page": 1, "Content": "No extractable text found in this PDF"}]
        df = pd.DataFrame(rows)
        df.to_excel(output_path, index=False)
    except Exception as e:
        raise RuntimeError(f"Failed to extract data from PDF: {e}")
    return output_filename
